"""
Build the submission document for the ET AI Hackathon 2026, PS-7.

Produces BOTH formats from ONE definition of the content, so they cannot drift:
    outputs/Resilience_Graph_AI_Submission.docx   (python-docx)
    outputs/Resilience_Graph_AI_Submission.pdf    (HTML twin printed by headless Chrome)

Design constraints (from the submission brief):
  - single black colour throughout, no colour fills, no emoji
  - no em dashes and no en dashes anywhere in the text
  - every table ruled with 1px solid black borders
  - every number traceable to reports/metrics.json,
    reports/scaling_measurements.json, or a labelled external citation

    ./.venv/Scripts/python.exe -m scripts.make_submission_doc

Word is deliberately NOT used for the PDF: Word COM automation hangs on an
invisible first-run dialog in this environment. Chrome renders the HTML twin
instead, and build() is written once against a backend interface that both
renderers implement.
"""
from __future__ import annotations

import html
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC_ASSETS = ROOT / "reports" / "doc"
OUT_DOCX = ROOT / "outputs" / "Resilience_Graph_AI_Submission.docx"
OUT_HTML = ROOT / "outputs" / "Resilience_Graph_AI_Submission.html"
OUT_PDF = ROOT / "outputs" / "Resilience_Graph_AI_Submission.pdf"
CHROME = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")

BLACK = RGBColor(0, 0, 0)
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
CONTENT_W = Cm(17.0)

M = json.loads((ROOT / "reports" / "metrics.json").read_text(encoding="utf-8"))
SCALING = json.loads((ROOT / "reports" / "scaling_measurements.json").read_text(encoding="utf-8"))

E1C, E1L, E1U = M["engine1"]["cicids"], M["engine1"]["lanl"], M["engine1"]["unsw"]
E2P = M["engine2"]["predictor"]
E2E = M["engine2"]["embeddings"]


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


# ==========================================================================
# DOCX backend
# ==========================================================================
def _shd_clear(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tcPr.append(shd)


def _set_table_borders(table, sz: int = 8) -> None:
    """1pt solid black on every edge (w:sz is in eighths of a point)."""
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


ALIGN = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}


class DocxBackend:
    def __init__(self):
        self.texts = []
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name = BODY_FONT
        st.font.size = Pt(10.5)
        st.font.color.rgb = BLACK
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        for s in doc.sections:
            s.page_width, s.page_height = Cm(21.0), Cm(29.7)
            s.left_margin = s.right_margin = Cm(2.0)
            s.top_margin = s.bottom_margin = Cm(2.0)
        self.doc = doc

    def _run(self, run, size=10.5, bold=False, italic=False, mono=False):
        run.font.name = MONO_FONT if mono else BODY_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = BLACK
        rPr = run._element.get_or_add_rPr()
        rPr.get_or_add_rFonts().set(qn("w:eastAsia"), MONO_FONT if mono else BODY_FONT)

    def para(self, text="", size=10.5, bold=False, italic=False, mono=False,
             align=None, space_after=6, space_before=0, indent=None):
        self.texts.append(text)
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after, pf.space_before, pf.line_spacing = Pt(space_after), Pt(space_before), 1.15
        if align:
            p.alignment = ALIGN[align]
        if indent is not None:
            pf.left_indent = Cm(indent)
        if text:
            self._run(p.add_run(text), size, bold, italic, mono)
        return p

    def rich(self, parts, size=10.5, space_after=6, indent=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if indent is not None:
            p.paragraph_format.left_indent = Cm(indent)
        for text, bold, mono in parts:
            self.texts.append(text)
            self._run(p.add_run(text), size, bold=bold, mono=mono)
        return p

    def heading(self, text, level=1):
        self.texts.append(text)
        sizes = {1: 15, 2: 12.5, 3: 11}
        before = {1: 16, 2: 12, 3: 9}
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before[level])
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True
        self._run(p.add_run(text), sizes[level], bold=True)
        return p

    def bullets(self, items, size=10.5):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.12
            if isinstance(it, list):
                for text, bold, mono in it:
                    self.texts.append(text)
                    self._run(p.add_run(text), size, bold=bold, mono=mono)
            else:
                self.texts.append(it)
                self._run(p.add_run(it), size)

    def table(self, rows, widths=None, size=9.5, header=True, mono_cols=()):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        _set_table_borders(t)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                self.texts.append(str(val))
                cell = t.cell(r, c)
                _shd_clear(cell)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                self._run(p.add_run(str(val)), size,
                          bold=(header and r == 0), mono=(c in mono_cols and r > 0))
        if widths:
            for c, w in enumerate(widths):
                for r in range(len(rows)):
                    t.cell(r, c).width = Cm(w)
        return t

    def code(self, lines):
        t = self.doc.add_table(rows=1, cols=1)
        _set_table_borders(t)
        cell = t.cell(0, 0)
        _shd_clear(cell)
        cell.text = ""
        cell.width = CONTENT_W
        for i, ln in enumerate(lines):
            self.texts.append(ln)
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            self._run(p.add_run(ln), 8.5, mono=True)
        self.para("", space_after=4)

    def figure(self, path, caption, width_cm=17.0):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        self.para(caption, size=9, italic=True, align="center", space_after=10)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)


# ==========================================================================
# HTML backend (printed to PDF by headless Chrome)
# ==========================================================================
CSS = """
@page { size: A4; margin: 20mm; }
* { box-sizing: border-box; }
body { font-family: Calibri, "Segoe UI", Arial, sans-serif; font-size: 10.5pt;
       color: #000; line-height: 1.15; margin: 0; background: #fff; }
p { margin: 0 0 6pt 0; }
h1 { font-size: 15pt; margin: 16pt 0 5pt 0; page-break-after: avoid; }
h2 { font-size: 12.5pt; margin: 12pt 0 5pt 0; page-break-after: avoid; }
h3 { font-size: 11pt; margin: 9pt 0 5pt 0; page-break-after: avoid; }
ul { margin: 0 0 6pt 0; padding-left: 16pt; }
li { margin-bottom: 3pt; line-height: 1.12; }
table { border-collapse: collapse; width: 100%; margin: 0 0 6pt 0;
        page-break-inside: avoid; }
td { border: 1px solid #000; padding: 3pt 5pt; font-size: 9.5pt;
     vertical-align: top; background: #fff; }
tr:first-child td { font-weight: bold; }
table.plain tr:first-child td { font-weight: normal; }
table.code tr:first-child td { font-weight: normal; font-family: Consolas, monospace;
     font-size: 8.5pt; white-space: pre; line-height: 1.0; }
figure { margin: 6pt 0 10pt 0; text-align: center; page-break-inside: avoid; }
figure img { width: 100%; }
figcaption { font-size: 9pt; font-style: italic; margin-top: 3pt; }
.pb { page-break-before: always; }
"""


class HtmlBackend:
    def __init__(self):
        self.texts = []
        self.parts = []
        self._pb = False

    def _cls(self, extra=""):
        names = ([extra] if extra else []) + (["pb"] if self._pb else [])
        self._pb = False
        return (" class='" + " ".join(names) + "'") if names else ""

    def para(self, text="", size=10.5, bold=False, italic=False, mono=False,
             align=None, space_after=6, space_before=0, indent=None):
        self.texts.append(text)
        st = ["font-size:%gpt" % size,
              "margin-bottom:%gpt" % space_after,
              "margin-top:%gpt" % space_before]
        if bold:
            st.append("font-weight:bold")
        if italic:
            st.append("font-style:italic")
        if mono:
            st.append("font-family:Consolas,monospace")
        if align:
            st.append("text-align:" + align)
        if indent is not None:
            st.append("margin-left:%gcm" % indent)
        self.parts.append("<p%s style='%s'>%s</p>"
                          % (self._cls(), ";".join(st), html.escape(text)))

    def rich(self, parts, size=10.5, space_after=6, indent=None):
        inner = ""
        for text, bold, mono in parts:
            self.texts.append(text)
            s = []
            if bold:
                s.append("font-weight:bold")
            if mono:
                s.append("font-family:Consolas,monospace")
            inner += "<span style='%s'>%s</span>" % (";".join(s), html.escape(text))
        st = "font-size:%gpt;margin-bottom:%gpt" % (size, space_after)
        if indent is not None:
            st += ";margin-left:%gcm" % indent
        self.parts.append("<p%s style='%s'>%s</p>" % (self._cls(), st, inner))

    def heading(self, text, level=1):
        self.texts.append(text)
        self.parts.append("<h%d%s>%s</h%d>" % (level, self._cls(), html.escape(text), level))

    def bullets(self, items, size=10.5):
        lis = ""
        for it in items:
            if isinstance(it, list):
                inner = ""
                for text, bold, mono in it:
                    self.texts.append(text)
                    s = []
                    if bold:
                        s.append("font-weight:bold")
                    if mono:
                        s.append("font-family:Consolas,monospace")
                    inner += "<span style='%s'>%s</span>" % (";".join(s), html.escape(text))
                lis += "<li style='font-size:%gpt'>%s</li>" % (size, inner)
            else:
                self.texts.append(it)
                lis += "<li style='font-size:%gpt'>%s</li>" % (size, html.escape(it))
        self.parts.append("<ul%s>%s</ul>" % (self._cls(), lis))

    def table(self, rows, widths=None, size=9.5, header=True, mono_cols=()):
        total = float(sum(widths)) if widths else None
        out = "<table%s>" % self._cls("" if header else "plain")
        for r, row in enumerate(rows):
            out += "<tr>"
            for c, val in enumerate(row):
                self.texts.append(str(val))
                st = ["font-size:%gpt" % size]
                if widths:
                    st.append("width:%.1f%%" % (widths[c] / total * 100.0))
                if c in mono_cols and r > 0:
                    st.append("font-family:Consolas,monospace")
                out += "<td style='%s'>%s</td>" % (";".join(st), html.escape(str(val)))
            out += "</tr>"
        self.parts.append(out + "</table>")

    def code(self, lines):
        for ln in lines:
            self.texts.append(ln)
        body = html.escape("\n".join(lines))
        self.parts.append("<table%s><tr><td>%s</td></tr></table>"
                          % (self._cls("code"), body))

    def figure(self, path, caption, width_cm=17.0):
        self.texts.append(caption)
        self.parts.append(
            "<figure%s><img style='width:%gcm' src='%s'>"
            "<figcaption>%s</figcaption></figure>"
            % (self._cls(), width_cm, Path(path).resolve().as_uri(), html.escape(caption)))

    def page_break(self):
        self._pb = True

    def save(self, path):
        doc = ("<!doctype html><html><head><meta charset='utf-8'>"
               "<title>Resilience Graph AI Submission</title>"
               "<style>" + CSS + "</style></head><body>"
               + "".join(self.parts) + "</body></html>")
        Path(path).write_text(doc, encoding="utf-8")


# ==========================================================================
# The document content, written once against the backend interface above.
# ==========================================================================
def build(B):
    # ---------------- cover ----------------
    B.para("", space_after=40)
    B.para("RESILIENCE GRAPH AI", size=26, bold=True, align="center",
         space_after=6)
    B.para("Behavioural detection, attack path reconstruction and guided response "
              "for critical national infrastructure",
         size=13, align="center", space_after=22)
    B.para("Problem Statement 7", size=12, bold=True,
         align="center", space_after=2)
    B.para("AI Driven Cyber Resilience for Critical National Infrastructure",
         size=11, align="center", space_after=26)
    B.para("ET AI Hackathon 2026", size=11.5, bold=True,
         align="center", space_after=2)
    B.para("Team rishiikumarsingh2201", size=11,
         align="center", space_after=20)

    B.table([
        ["Team members", "Rishii Kumar Singh, Sarthak Tomar, Aman Kumar, Aarushi Aanand"],
    ], widths=[4.4, 12.6], size=10, header=False)
    B.para("", space_after=14)

    B.para("Project links", size=10.5, bold=True, space_after=4)
    B.table([
        ["Resource", "Link"],
        ["Live application (Render)", "https://resilience-graph-ai.onrender.com"],
        ["Source code (GitHub)", "https://github.com/RishiiGamer2201/resilience-graph-ai"],
        ["Presentation (Canva)", "https://canva.link/f4gesmsduelihuz"],
        ["Submission document", "https://docs.google.com/document/d/"
                                "10fs9PBoaEQLUJkyoJPSqgNp-ejRsWS90C37CWV8kDgU/edit"],
        ["Dataset", "https://www.kaggle.com/datasets/c3c7d72d2098d35857c2136a6d1c357"
                    "85b7ba94e0f48ed6de68d0ab1ed021945"],
        ["Competition", "https://unstop.com/competitions/"
                        "crp-et-ai-hackathon-20-economic-times-1675680"],
    ], widths=[4.4, 12.6], size=8.5)

    B.para("", space_after=10)
    B.para("A note on the numbers in this document: every figure is measured, and each one "
              "traces to an evaluation report or a measurement file in the repository. Where a "
              "number comes from outside our own work it is attributed to its source. Where a "
              "result is weak, or where a component is simulated rather than live, this document "
              "says so directly.",
         size=9.5, italic=True)

    B.page_break()

    # ---------------- contents ----------------
    B.heading("Contents", 1)
    B.table([
        ["Section", "Title"],
        ["1", "Executive summary"],
        ["2", "The problem"],
        ["3", "Our solution"],
        ["4", "How the system works, end to end"],
        ["5", "System architecture"],
        ["6", "Engine 1: behavioural detection"],
        ["7", "The shared spine: turning alerts into a story"],
        ["8", "Engine 2: prediction and attribution"],
        ["9", "Threat Radar: external intelligence"],
        ["10", "The application"],
        ["11", "Data: what we used and why"],
        ["12", "Technology choices, and what we rejected"],
        ["13", "Results in full"],
        ["14", "Performance and scalability"],
        ["15", "How we kept ourselves honest"],
        ["16", "Testing and engineering"],
        ["17", "Limitations we state before you find them"],
        ["18", "Impact, deployment path and roadmap"],
        ["19", "Reproducing this work"],
        ["20", "Glossary"],
    ], widths=[2.0, 15.0], size=10)

    B.page_break()

    # ---------------- 1 executive summary ----------------
    B.heading("1. Executive summary", 1)
    B.para("An attacker breaks into a hospital network. They do not break anything. They "
              "steal one employee's password and quietly log into machine after machine, looking "
              "for the patient database. Every one of those logins looks completely normal on its "
              "own. That is why intrusions of this kind go undetected for a global median of "
              "about ten days.")
    B.para("Resilience Graph AI reads the ordinary authentication logs an organisation already "
              "collects, and finds the story hidden across them. It does not look for known bad "
              "software. It learns what normal behaviour looks like for each account, then "
              "measures deviation from it.")
    B.para("On a real, labelled red team campaign the system produces the following, live, in "
              "under one fifth of a second:", space_after=6)

    B.table([
        ["What the system does", "Measured result"],
        ["Scores every authentication event and flags the anomalies",
         "2,732 events scored, 1,192 flagged"],
        ["Collapses those alerts into a single narrated incident",
         "1,192 alerts become 1 incident"],
        ["Reconstructs the attacker's movement across the estate",
         "479 machines, 502 movements, 4 attacker controlled hosts"],
        ["Identifies the single best machine to isolate first",
         "Isolating 1 host severs 463 machines of exposure"],
        ["Detects the attack without ever seeing an attack label in training",
         f"ROC-AUC {E1L['roc_auc']} against 702 real red team events"],
    ], widths=[9.0, 8.0])
    B.para("", space_after=8)

    B.para("The system additionally maps every step to the MITRE ATT&CK catalogue, the "
              "industry standard naming scheme for attacker techniques, predicts the likely next "
              "technique with a real transition probability, ranks which known threat group the "
              "behaviour resembles with a written justification, and recommends containment that "
              "a human must approve before anything happens.")
    B.para("Three properties separate this from a demonstration. First, the whole pipeline "
              "runs live on any log submitted to it, including a file uploaded by a judge. "
              "Second, no number displayed anywhere in the application is hard coded; each one "
              "is computed by the analysis that is currently loaded. Third, the parts that are "
              "simulated, specifically the containment actions, are labelled as simulated "
              "everywhere they appear.")

    # ---------------- 2 problem ----------------
    B.heading("2. The problem", 1)

    B.heading("2.1 The Indian context", 2)
    B.table([
        ["Fact", "Figure", "Source"],
        ["Cyber security incidents handled by CERT-In during 2023", "1.59 million and above",
         "CERT-In"],
        ["Indian government entities operating end of life IT", "Over 70 percent",
         "PS-7 problem brief"],
        ["Global median attacker dwell time", "About 10 days", "Mandiant M-Trends 2024"],
    ], widths=[8.2, 4.6, 4.2])
    B.para("", space_after=8)
    B.para("Two Indian precedents shaped this project directly. The ransomware attack on AIIMS "
              "Delhi in 2022 took the country's premier hospital offline for days. The breaches "
              "at the Central Board of Secondary Education affected the national school "
              "examination system. In both categories of organisation the defining constraint is "
              "the same: critical systems, legacy infrastructure, and no capacity to staff a "
              "round the clock security operations centre.")

    B.heading("2.2 Why existing defences miss this class of attack", 2)
    B.table([
        ["Defence", "How it works", "Why it misses a quiet credential based intrusion"],
        ["Antivirus and signatures", "Matches known bad files and patterns",
         "The attacker deploys no malware. They use valid logins."],
        ["Firewall", "Blocks unauthorised network connections",
         "The attacker is already inside, using permitted paths."],
        ["Threshold rules, for example alert on five failed logins",
         "A limit on a single counter",
         "A patient attacker simply stays under the threshold."],
        ["Standard SIEM alerting", "Scores each event independently",
         "One intrusion becomes thousands of disconnected alerts, and the pattern across them "
         "is never assembled."],
    ], widths=[3.6, 5.2, 8.2])
    B.para("", space_after=8)

    B.heading("2.3 The three failures we set out to fix", 2)
    B.bullets([
        [("Signature evasion. ", True, False),
         ("An attack carried out with stolen but valid credentials matches no known bad rule, "
          "because nothing about it is technically unauthorised.", False, False)],
        [("Alert fatigue. ", True, False),
         ("Because each event is scored alone, a single intrusion is presented to an analyst as "
          "thousands of separate rows. Analysts triage rows, miss the pattern, and burn out.",
          False, False)],
        [("No blast radius view. ", True, False),
         ("Even when one alert is investigated properly, nobody can see the path from a "
          "compromised workstation to the patient database, or answer the only question that "
          "matters during an incident: which single machine do we unplug first.", False, False)],
    ])
    B.para("The insight this project rests on is that the data needed to catch these attacks "
              "already exists, in logs organisations already collect and already pay to store. "
              "The missing component is the layer that connects those records to each other. "
              "That is what we built, and it is why the system needs no new sensors, no new "
              "agents on endpoints, and no change to existing infrastructure.")

    B.page_break()

    # ---------------- 3 solution ----------------
    B.heading("3. Our solution", 1)
    B.para("Resilience Graph AI is a working web application backed by two AI engines joined "
              "by a shared analysis pipeline. It accepts an authentication log, either one of the "
              "scenarios shipped with the product or a file uploaded by the user, and returns a "
              "complete incident investigation.")
    B.table([
        ["Capability", "What the user gets", "Status"],
        ["Analyse any log", "Select a scenario or upload a CSV. Every screen in the application "
                            "re-renders on that data.", "Live, per request"],
        ["Campaign view", "All 104 compromised accounts presented as one campaign rather than a "
                          "single victim.", "Live"],
        ["Per account investigation", "Open any account to get its own scoped incident, graph "
                                      "and report.", "Live"],
        ["Attack path graph", "Click any machine to see every authentication involving it, and "
                              "the blast radius across all attacker footholds.", "Live"],
        ["Single event scoring", "Score one authentication event on demand using the real "
                                 "trained model.", "Live"],
        ["Next technique prediction", "A ranked next move with a genuine transition probability.",
         "Live"],
        ["Threat group attribution", "A ranked MITRE group with an auditable written "
                                     "justification.", "Live, transparent retrieval"],
        ["Threat Radar", "External threat intelligence, India first, mapped to ATT&CK and cross "
                         "referenced against the current incident.", "Live feeds"],
        ["Audit ready report", "A printable incident report suitable for compliance records.",
         "Live"],
        ["Guided containment", "Recommended response steps. Anything touching a critical asset "
                               "requires human approval.", "Simulated by design"],
    ], widths=[4.0, 9.6, 3.4])
    B.para("", space_after=8)
    B.para("The final row is deliberate. There is no live production network attached to this "
              "system, so no containment action is actually executed. Every such action is "
              "labelled as simulated in the interface. We consider stating this plainly to be "
              "part of the engineering, not a caveat to it.")

    # ---------------- 4 pipeline ----------------
    B.heading("4. How the system works, end to end", 1)
    B.para("The analysis is a single function call that executes seven stages in order. "
              "Stages one and two are Engine 1, which finds the anomalies. Stages three to six "
              "are the shared spine, which turns those anomalies into an incident a human can "
              "act on. Stage seven is Engine 2, which looks forward and outward.")
    B.figure(DOC_ASSETS / "pipeline_bw.png",
           "Figure 1. The seven stage analysis pipeline. Every stage executes on each request.")

    B.page_break()

    # ---------------- 5 architecture ----------------
    B.heading("5. System architecture", 1)
    B.para("The application is one repository and deploys as one container. The frontend is a "
              "React single page application. The backend is FastAPI, which serves both the "
              "programming interface and, in production, the built frontend from the same origin. "
              "The analysis spine sits behind the live endpoints. Trained models and lookup "
              "tables are loaded once at start up.")
    B.figure(ROOT / "reports" / "technical_architecture_final.png",
             "Figure 2. Technical architecture. Inputs, the single container runtime across four "
             "planes, and the outcomes delivered to a security team. A full resolution copy is in "
             "the repository at reports/technical_architecture_final.png.",
             width_cm=17.0)

    B.heading("5.1 The programming interface", 2)
    B.table([
        ["Type", "Endpoints", "What happens"],
        ["Live analysis", "POST /api/analyze, POST /api/analyze/upload",
         "Runs the complete seven stage pipeline on the submitted log and returns one bundle "
         "containing the payload for every screen."],
        ["Live model calls", "POST /api/score-event, POST /api/predict-next",
         "Single shot calls into the trained anomaly model and the transition model."],
        ["Live intelligence", "POST /api/threat-radar",
         "Fetches external threat intelligence and scores it against the current incident."],
        ["Streaming", "GET /api/analyze/stream",
         "Server sent events, replaying an incident event by event."],
        ["Cached reads", "GET /api/overview, /incident, /graph, /report, /threat-intel, "
                         "/attackers, /metrics",
         "Serve a sample that is itself a real analysis of a shipped log, produced by calling "
         "the same pipeline offline."],
        ["Static", "GET /", "Serves the built frontend in production."],
    ], widths=[3.0, 5.4, 8.6], size=9)
    B.para("", space_after=8)
    B.para("The last two rows matter for a specific reason. There is no separate mock data "
              "path in this system. The sample content that loads before a user runs their own "
              "analysis was generated by running the real pipeline over a real log and saving "
              "the result. A demonstration path that could quietly diverge from the real one "
              "does not exist, because it was never built.")

    B.heading("5.2 Deployment topology", 2)
    B.table([
        ["Environment", "Topology"],
        ["Local development", "Two processes. The backend runs on port 8000, the frontend "
                              "development server on port 5173 and proxies interface calls to it."],
        ["Production", "One Docker container built in two stages. A Node stage builds the "
                       "frontend, a slim Python stage serves the interface and the built "
                       "frontend from a single origin, so no cross origin configuration is "
                       "required. Hosted on Render."],
        ["Runtime requirements", "No GPU. The deep learning framework is not installed in the "
                                 "deployed image at all. Sentence embeddings ship as a "
                                 "precomputed file."],
        ["Cold start from a clone", "Trained models, ATT&CK lookups, embeddings, demonstration "
                                    "scenarios and the sample cache are all committed, so the "
                                    "application runs from a fresh clone with no dataset "
                                    "download."],
    ], widths=[4.2, 12.8])

    B.page_break()

    # ---------------- 6 engine 1 ----------------
    B.heading("6. Engine 1: behavioural detection", 1)

    B.heading("6.1 The idea", 2)
    B.para("We never tell the model what an attack looks like. We show it a large volume of "
              "normal behaviour and ask a single question of each new event: how unusual is this. "
              "This is the only approach that can catch an attack nobody has catalogued yet, "
              "because it does not require the attack to have been seen before.")
    B.para("The difficulty is choosing what to measure. Raw log fields such as a user name or "
              "a machine name are useless to a model, because they are arbitrary identifiers. So "
              "we compute behavioural features that describe how a person moves through a "
              "network, and we compute them chronologically, per account.")

    B.heading("6.2 The seven behavioural features", 2)
    B.table([
        ["Feature", "What it measures", "Why an attacker trips it"],
        ["new_dst_for_user", "First time this account has ever logged into this machine",
         "Attackers explore machines the real user never touches"],
        ["new_src_for_user", "First time this account logged in from this machine",
         "The attacker operates from their own foothold machine"],
        ["user_distinct_dst_sofar", "How many different machines this account has reached so far",
         "Rapid fan out is hunting behaviour, not daily work"],
        ["user_fail_rate_sofar", "Running share of failed logins for this account",
         "Credential guessing leaves a trail of failures"],
        ["dst_rarity", "How rarely anyone at all logs into this destination",
         "Attackers reach obscure but high value servers"],
        ["is_fail", "Whether this login failed", "Basic corroborating signal"],
        ["is_ntlm", "Whether an older authentication protocol was used",
         "Pass the hash attacks depend on it"],
    ], widths=[4.4, 6.2, 6.4], size=9, mono_cols=(0,))
    B.para("", space_after=8)
    B.para("These are not chosen by intuition alone. Measured on the real labelled data, the "
              "separation between benign and attacker behaviour is as follows.", space_after=6)
    B.table([
        ["Feature", "Benign average", "Attacker average", "Separation"],
        ["new_dst_for_user", "0.0204", "0.2821", "13.8 times"],
        ["new_src_for_user", "0.0174", "0.1054", "6.1 times"],
        ["user_fail_rate_sofar", "0.0015", "0.0081", "5.4 times"],
        ["dst_rarity", "4.8715", "9.8256", "2.0 times"],
        ["is_ntlm", "0.0580", "1.0000", "Present in 100 percent of attack events"],
    ], widths=[4.4, 3.4, 3.4, 5.8], size=9, mono_cols=(0,))
    B.para("", space_after=8)
    B.para("An attacker is 13.8 times more likely to touch a machine that the account they "
              "stole has never visited. That single behavioural fact carries most of the "
              "detection.")

    B.heading("6.3 The model and its exact configuration", 2)
    B.para("The shipped detector is an Isolation Forest. It isolates outliers by splitting the "
              "data at random. Anomalies are isolated in fewer splits because they sit apart from "
              "the crowd. It is fast, it requires no labels, and it behaves well in higher "
              "dimensions.")
    B.table([
        ["Setting", "Value and reasoning"],
        ["Estimators", "200 trees"],
        ["Maximum samples", "4,096 per tree"],
        ["Contamination", "Set automatically"],
        ["Random seed", "42, fixed so results reproduce exactly"],
        ["Training data", "A random sample of 800,000 rows labelled benign. Attack rows are "
                          "excluded from training entirely."],
        ["Normalisation", "Standard scaling, fitted on the training sample only, so no "
                          "information from the evaluation data reaches the model."],
        ["Score calibration", "Raw scores are mapped to a 0 to 100 range using fixed anchor "
                              "vectors stored in the repository, not the minimum and maximum of "
                              "the current batch. A score therefore means the same thing across "
                              "different uploads, and matches the single event endpoint exactly."],
        ["Use of labels", "Evaluation only. Labels never enter training. This is what makes the "
                          "reported detection score defensible rather than circular."],
    ], widths=[4.2, 12.8], size=9.5)
    B.para("", space_after=8)
    B.para("There is no look ahead leakage. Every running statistic, such as the count of "
              "distinct machines an account has reached, is computed using only that account's "
              "prior events, so the score of any given event never depends on the future.")

    B.page_break()

    # ---------------- 7 spine ----------------
    B.heading("7. The shared spine: turning alerts into a story", 1)
    B.para("This is the part of the system that converts a pile of anomalies into something a "
              "human being can act on within minutes.")

    B.heading("7.1 Correlation: 1,192 alerts become 1 incident", 2)
    B.para("Any event scoring 50 or above becomes an alert. Rather than emitting 1,192 "
              "separate alerts, the system groups them into one incident carrying a timeline, a "
              "severity, the list of affected accounts, and an ordered chain of techniques that "
              "reads as a narrative. An hour of silence starts a new session. Severity is taken "
              "from the highest scoring event in the incident: 90 and above is critical, 75 and "
              "above is high, 50 and above is medium.")

    B.heading("7.2 Mapping behaviour to the MITRE ATT&CK catalogue", 2)
    B.para("MITRE ATT&CK is a free public catalogue that gives every known attacker technique "
              "an identifier, and it is the common vocabulary of the security industry. We infer "
              "the technique from observed behaviour, never from a text label.")
    B.table([
        ["Observed behaviour", "Mapped technique", "Plain meaning"],
        ["The login failed", "T1110", "Brute force, guessing passwords"],
        ["New machine reached using the older protocol", "T1550.002",
         "Pass the hash, reusing a stolen credential fingerprint"],
        ["New machine reached", "T1021", "Remote services, legitimate remote login tools"],
        ["Neither condition met", "No technique", "Treated as normal activity"],
    ], widths=[5.6, 3.4, 8.0], size=9.5, mono_cols=(1,))
    B.para("", space_after=8)
    B.para("Technique names, descriptions and recommended mitigations are read from the "
              "official MITRE data files, which we parse ourselves into a lookup table covering "
              "918 techniques and 175 groups. The explanation text shown to the user is MITRE's "
              "own wording. No language model generates technique text anywhere in this system, "
              "which means a fabricated technique identifier is structurally impossible rather "
              "than merely unlikely.")

    B.heading("7.3 The attack path graph", 2)
    B.para("Every alert becomes an edge in a directed graph running from source machine to "
              "destination machine. From that graph we compute the four things a responder "
              "actually needs.")
    B.table([
        ["Question a responder asks", "How the graph answers it", "Result on the real campaign"],
        ["Where is the attacker operating from", "Machines with outbound attack movement",
         "4 footholds. One machine, C17693, carries 670 of the 702 red team events."],
        ["How far can they reach", "Reachable set from every foothold, combined",
         "475 machines"],
        ["Which valuable machines are exposed", "Shortest path to each critical asset",
         "18 critical assets reachable"],
        ["What do we disconnect first", "Ranking by betweenness centrality",
         "Isolating C17693 alone severs 463 machines"],
    ], widths=[4.4, 5.2, 7.4], size=9)
    B.para("", space_after=8)
    B.para("We deliberately report two different numbers rather than one flattering one. Total "
              "exposure is 475 machines. What isolating a single choke point actually severs is "
              "463. Presenting only the larger figure would overstate what one containment action "
              "achieves.")
    B.para("This distinction came out of a real defect we found and fixed. An earlier version "
              "assumed the attacker had one entry point, computed reachability from that machine "
              "alone, and consequently under reported exposure and wrongly cleared four critical "
              "assets as safe. The corrected version takes the union of reachable sets over every "
              "foothold, and a regression test now fails if anyone reintroduces the assumption.")
    B.para("On critical assets, an honest note. The public dataset is anonymised and carries "
              "no asset criticality labels at all. We therefore derive them from a stated "
              "heuristic: the machines that the largest number of distinct accounts authenticate "
              "to, which in practice surfaces domain controllers and authentication servers. On "
              "that basis the red team reached 13 of the estate's 20 most depended upon servers, "
              "including one that 17,808 accounts rely on. In a real deployment the operator "
              "supplies their own asset list instead. It is already an input parameter, not a "
              "value written into the code.")

    B.heading("7.4 Guided response, with a human in the loop", 2)
    B.table([
        ["Situation", "Response mode"],
        ["A critical asset is involved", "Requires explicit human approval"],
        ["High or critical severity", "Simulated containment"],
        ["Medium severity", "Raise a ticket and enrich it"],
        ["Low severity", "Monitor only"],
    ], widths=[8.5, 8.5])
    B.para("", space_after=8)
    B.para("Recommended containment steps are seeded from the real MITRE mitigations "
              "associated with the observed techniques, so the advice is MITRE's rather than "
              "ours. Every action is simulated, and labelled as such wherever it appears.")

    B.page_break()

    # ---------------- 8 engine 2 ----------------
    B.heading("8. Engine 2: prediction and attribution", 1)

    B.heading("8.1 Predicting the attacker's next move", 2)
    B.para("Given the techniques observed so far, what is likely to come next. We learned "
              "technique to technique transitions from 201 real attack sequences, made up of 145 "
              "MITRE group profiles and 56 campaigns, each containing at least six techniques, "
              "plus 4 analyst verified CERT-In advisories, giving 205 sequences in total.")
    B.table([
        ["Method", "Top 3 accuracy", "Outcome"],
        ["Most frequent technique baseline", pct(E2P["most_frequent_top3"]), "Baseline"],
        ["Kill chain order baseline", pct(E2P["killchain_top3"]),
         "Baseline built specifically to beat us"],
        ["Recurrent neural network over sentence embeddings", pct(E2P["lstm_top3"]),
         "Lost. Published as a documented negative result."],
        ["First order Markov chain", pct(E2P["markov_top3"]), "Shipped"],
    ], widths=[6.4, 3.2, 7.4], size=9.5)
    B.para("", space_after=8)

    B.para("Two things here matter more than the headline number.", bold=False)
    B.rich([("The circularity trap, and how we escaped it. ", True, False),
               ("Our sequences are ordered using MITRE's kill chain tactic order, which runs from "
                "reconnaissance through to impact. A model could score well simply by re-learning "
                "that ordering, while learning nothing whatsoever about attacks. So we built a "
                "baseline whose entire strategy is to exploit that ordering, and we required our "
                "model to beat it. The Markov model beats it by a factor of 5.2, which is "
                "evidence that it is predicting genuine technique to technique transitions "
                "rather than re-deriving a sort order.", False, False)])
    B.rich([("We shipped the model that won, not the impressive one. ", True, False),
               (f"The neural network scored {pct(E2P['lstm_top3'])} and lost to a first order "
                "Markov chain at this data scale. We ship the Markov model and publish the "
                "neural result as a negative finding. At 205 sequences, a simple model is the "
                "correct engineering answer, and reporting the loss is more useful to a reader "
                "than hiding it.", False, False)])
    B.para("The prediction returned by the interface is a genuine first order transition "
              "probability, computed as the observed count of a transition divided by the total "
              "from that state. An example from the shipped model is a transition observed at "
              "52.5 percent. Where a state has never been seen, the system falls back to a "
              "frequency ranked list and explicitly scores those suggestions at zero, because "
              "there is no observed evidence for them. Data splits are made at the level of whole "
              "sequences, never inside one, which prevents leakage.")
    B.rich([("The non circular India test. ", True, False),
               ("The four CERT-In sequences are ordered by the real reported timeline rather than "
                "by our heuristic, and they appear only in the test set. The model scores "
                f"{pct(M['engine2']['manual_cert_in_top3'])} top 3 on them against "
                f"{pct(E2P['markov_top3'])} on the automatically ordered set. We treat that gap "
                "as a finding rather than a failure: real attacker orderings are harder, which "
                "demonstrates that part of the higher number was ordering driven. Both figures "
                "are published.", False, False)])

    B.heading("8.2 Attributing the activity to a known threat group", 2)
    B.para("Given the observed techniques, which publicly documented threat group does this "
              "resemble. We score the observed technique set against 172 MITRE group profiles "
              "using a transparent weighted formula rather than a trained classifier.")
    B.code([
        "score = 0.55 * coverage  +  0.20 * jaccard  +  0.25 * semantic_similarity",
        "",
        "coverage             fraction of observed techniques in the group's public profile",
        "jaccard              set overlap between observed and profile techniques",
        "semantic_similarity  cosine distance between sentence embedding centroids",
    ])
    B.para("Every result carries a generated justification, for example that it matches three "
              "of four observed techniques, with a stated profile coverage and semantic "
              "similarity. A user can therefore audit the reasoning rather than trust a score.")
    B.rich([("The caveat we state before anyone asks. ", True, False),
               ("The built in evaluation, which hides 40 percent of a group's profile and then "
                "retrieves the group from the remainder, scores 100 percent at rank one. That "
                "number is close to trivial by construction, because it retrieves a public "
                "profile from a piece of itself. We never present it as a headline result. This "
                "component is transparent retrieval with a printed rationale, not a trained "
                "classifier, and on an authentication only log carrying three techniques the "
                "ranked group is context for an analyst, not a conclusion.", False, False)])

    # ---------------- 9 radar ----------------
    B.heading("9. Threat Radar: external intelligence", 1)
    B.para("A security team also needs to know what is happening outside its own walls. Threat "
              "Radar pulls free and legitimate threat intelligence feeds, maps each item to real "
              "ATT&CK identifiers, and cross references those against the incident currently "
              "loaded. Items relevant to India are ranked first, because the problem statement "
              "concerns Indian infrastructure. That covers mentions of India, of Indian agencies "
              "such as CERT-In and NCIIPC, of Indian sectors such as UPI, Aadhaar and the "
              "Reserve Bank, and of actors known to target India. At the time of writing that is "
              "10 of 40 items.")
    B.para("The implementation uses only the Python standard library for network and document "
              "handling, so the deployed image gains no new dependencies. Each feed is isolated, "
              "meaning a dead source reports its own failure and never breaks the radar.")
    B.para("Technique mapping is deliberately precision first. It accepts explicit technique "
              "identifiers, then a curated table of about 70 phrases, then technique name "
              "matching. Reconnaissance and resource development techniques are excluded from "
              "name matching, because their official names are ordinary English nouns such as "
              "Software and Vulnerabilities, which matched innocent prose during testing. Every "
              "identifier is validated against the real catalogue before it is displayed.")
    B.para("Three honest notes about this feature.", space_after=4)
    B.bullets([
        [("CERT-In has no working feed. ", True, False),
         ("Their address returns a success code together with an HTML page saying the address "
          "was not found, which is a disguised failure. We detected it, excluded the source, and "
          "added a guard that rejects HTML pretending to be a feed.", False, False)],
        [("We do not scrape social media. ", True, False),
         ("This was evaluated and rejected. It violates platform terms, it is actively blocked, "
          "and attributing activity to named individuals from public posts risks accusing the "
          "wrong people. We consider that an unacceptable failure mode for a security product.",
          False, False)],
        [("No matches is a legitimate result. ", True, False),
         ("Our demonstration incident is authentication based while the public news cycle is "
          "dominated by software vulnerabilities and malware, so the honest answer is often that "
          "there are no matches. The interface says exactly that, rather than manufacturing a "
          "connection to look impressive.", False, False)],
    ])

    B.page_break()

    # ---------------- 10 application ----------------
    B.heading("10. The application", 1)
    B.para("Eight screens, all rendering the output of the analysis that is currently loaded.")
    B.table([
        ["Screen", "What it shows"],
        ["Analyze Log", "Choose a shipped scenario or upload a CSV, then run the full pipeline."],
        ["Overview", "Time to first alert, the active incident, and detector benchmarks."],
        ["Attackers", "All 104 compromised accounts. Open any one for its own scoped incident."],
        ["Live Incident", "Event by event replay, single event scoring, and the audit ready "
                          "report."],
        ["Attack Graph", "The interactive machine graph. Click any machine, or filter by account."],
        ["Threat Intel", "ATT&CK mapping, ranked threat groups, and live next technique "
                         "prediction."],
        ["Threat Radar", "External intelligence cross referenced with the current incident."],
        ["Models and Metrics, Data and Methodology", "The evidence tables, dataset descriptions "
                                                     "and honesty notes."],
    ], widths=[5.0, 12.0], size=9.5)
    B.para("", space_after=8)
    B.rich([("The single most important element in the interface ", True, False),
               ("is the badge in the top bar, which reads either LIVE ANALYSIS or SAMPLE DATA. "
                "A viewer can always tell whether what they are looking at was computed moments "
                "ago from their own data, or is the shipped sample. We consider that badge a "
                "correctness feature rather than a decoration.", False, False)])

    B.heading("10.1 Robustness for real uploads", 2)
    B.para("Two pieces of defensive engineering exist because of real failures, not "
              "speculation. The schema layer resolves column aliases case insensitively, so a "
              "column named username, account, src, source, dst or proto is understood without "
              "the user renaming anything. Timestamps are accepted either as epoch integers or "
              "as ISO-8601 date strings.")
    B.para("The second of those came from automated browser testing, which uploaded a file "
              "with ISO-8601 timestamps and found a crash, because every one of our own test "
              "files happened to use epoch integers. Real logs use dates. It is fixed, with a "
              "regression test that fails if it returns.")

    # ---------------- 11 data ----------------
    B.heading("11. Data: what we used and why", 1)
    B.table([
        ["Dataset", "Size", "Why it is in this project"],
        ["LANL Cyber Security Events", "11.2 million authentication events, 702 labelled red "
                                       "team events",
         "The core asset. A real attack campaign inside a real network, with the attacker's own "
         "events labelled, which lets us prove detection rather than assert it."],
        ["CIC-IDS2017", "2.3 million network flows",
         "A network level benchmark containing distinct attack families, used to compare model "
         "families fairly."],
        ["UNSW-NB15", "175,000 train and 82,000 test, official split",
         "A second, independent benchmark, included to show the approach generalises beyond one "
         "dataset."],
        ["MITRE ATT&CK, Enterprise, ICS and Mobile", "918 techniques, 175 groups",
         "The technique vocabulary, the official mitigations, and the group profiles used for "
         "attribution."],
        ["CERT-In advisories", "4 analyst verified sequences",
         "Real Indian attack timelines, used as the non circular test set for prediction."],
    ], widths=[3.8, 4.0, 9.2], size=9)
    B.para("", space_after=8)

    B.heading("11.1 Why the LANL dataset matters most", 2)
    B.para("It is a real 58 day capture from Los Alamos National Laboratory that includes an "
              "actual red team exercise, and critically, the red team's own events are labelled. "
              "Most public security datasets are either synthetic or unlabelled. This one lets us "
              "state a detection figure that can be checked.")
    B.para("The raw authentication file is roughly 70 gigabytes uncompressed. Our preparation "
              "script streams the compressed file rather than expanding it, reading 519 million "
              "lines, keeping every event involving a compromised account plus a one in four "
              "hundred background sample, joining the red team labels on the combination of time, "
              "source account, source machine and destination machine, and stopping early once "
              "past the attack window. The result is 11.2 million rows containing 702 malicious "
              "events. We note openly that 98.2 percent of the 715 red team records have an exact "
              "authentication counterpart, a known quirk of the dataset that we state rather than "
              "quietly round away.")
    B.para("CIC-IDS2017 is split by day, training on Monday to Wednesday benign traffic only "
              "and testing on Thursday and Friday, which contain seven attack families the model "
              "never saw. The destination port column is dropped so the model cannot simply "
              "memorise identifiers. Both choices exist to prevent the model scoring well for the "
              "wrong reason.")

    B.heading("11.2 Why we added the Mobile catalogue mid project", 2)
    B.para("A teammate's verified CERT-In advisory described an Android banking trojan whose "
              "techniques did not exist in the Enterprise and ICS catalogues. Every one of those "
              "identifiers would have been silently discarded, and the sequence would have been "
              "quietly wrong. Adding the Mobile matrix took the catalogue from 794 to 918 "
              "techniques. This matters for the problem statement, because India's threat "
              "landscape is heavily mobile.")

    B.page_break()

    # ---------------- 12 tech choices ----------------
    B.heading("12. Technology choices, and what we rejected", 1)
    B.para("Every significant choice below was made against a named alternative. Where the "
              "alternative won on a particular benchmark, we say so.")

    B.heading("12.1 Modelling choices", 2)
    B.table([
        ["Decision", "What we chose", "What we rejected, and why"],
        ["Detection approach", "Unsupervised anomaly detection trained only on benign data",
         "Supervised classification. It needs labelled attacks, which real organisations do not "
         "have, and it can only recognise attack types present in its training set. Our whole "
         "problem is the attack nobody has catalogued."],
        ["Detection model", "Isolation Forest",
         "One class support vector machines, which scale poorly to millions of rows. A deep "
         "autoencoder was also built and did beat Isolation Forest on the network dataset, at "
         f"{E1C['autoencoder_prauc']:.3f} against {E1C['iforest_prauc']:.3f} precision recall area. We "
         "report that openly, and still ship Isolation Forest for the authentication task, where "
         "it performs strongly and needs no GPU."],
        ["Prediction model", "First order Markov chain",
         f"A recurrent neural network over sentence embeddings, which scored "
         f"{pct(E2P['lstm_top3'])} against the Markov model's {pct(E2P['markov_top3'])}. At 205 "
         "sequences the simpler model is genuinely better. We shipped the winner and published "
         "the loss."],
        ["Attribution method", "Transparent weighted retrieval with a printed rationale",
         "A trained classifier. With 172 groups and few labelled campaigns it would overfit, and "
         "it could not explain itself. An analyst must be able to audit an attribution claim."],
        ["Technique descriptions", "Read verbatim from the official MITRE data files",
         "Generating technique text with a language model. That introduces the possibility of a "
         "confidently invented technique identifier, which in a security tool is a serious "
         "defect, not a cosmetic one."],
        ["Evaluation metric", "Precision recall area under curve, and true positive rate at a "
                              "fixed false positive rate",
         "Accuracy. At an attack prevalence of 0.006 percent, a model that answers benign every "
         "single time scores 99.994 percent accuracy while catching nothing at all."],
    ], widths=[3.2, 4.6, 9.2], size=9)

    B.heading("12.2 Engineering choices", 2)
    B.table([
        ["Decision", "What we chose", "What we rejected, and why"],
        ["Graph analytics", "networkx, in memory",
         "A dedicated graph database such as Neo4j. It would add an entire service to operate "
         "for no benefit at our data scale, where analysis completes in seconds. We state the "
         "ceiling of roughly 50,000 events per analysis and name the graph database as the "
         "documented next step rather than pretending the limit does not exist."],
        ["Backend framework", "FastAPI with uvicorn",
         "Django, which brings an ORM and admin layer we have no use for. Flask, which would "
         "need extra parts for streaming and request validation that FastAPI provides directly."],
        ["Frontend", "React with Vite",
         "Next.js. Server side rendering earns nothing for an analyst tool behind a login, and "
         "it would complicate packaging the whole product as one container."],
        ["Threat intelligence client", "Python standard library for network and document parsing",
         "Third party request and feed parsing libraries. Avoiding them keeps the deployed image "
         "at zero additional dependencies, which matters on a constrained free hosting tier."],
        ["Sentence embeddings", "A pretrained compact sentence transformer, precomputed at build "
                                "time and shipped as a file",
         "Calling a hosted embedding service at runtime, which would add cost, latency, an "
         "external dependency and a privacy question. Our approach means the deep learning "
         "framework is not installed in the production image at all."],
        ["Deployment", "One Docker container on a free hosting tier",
         "A multi service architecture. For a single analyst tool it would add operational "
         "burden and cost with no user visible benefit. One container means one address and a "
         "reviewer can run it in minutes."],
        ["Log ingestion", "CSV upload and shipped scenarios",
         "A direct live connector to a security platform. We have no such platform to connect "
         "to, and building an untestable integration would be theatre. Connectors are the first "
         "item on the roadmap."],
        ["External data collection", "Free, legitimate, published threat feeds",
         "Scraping social media platforms. It breaches platform terms, it is actively blocked, "
         "and person level attribution from public posts risks naming innocent people."],
    ], widths=[3.2, 4.6, 9.2], size=9)

    B.page_break()

    # ---------------- 13 results ----------------
    B.heading("13. Results in full", 1)

    B.heading("13.1 Detection", 2)
    B.table([
        ["Dataset", "Metric", "Result"],
        ["LANL", "ROC area under curve", str(E1L["roc_auc"])],
        ["LANL", "True positive rate at 5 percent false positive rate",
         f"{pct(E1L['tpr_at_5pct_fpr'])} (680 of 702 attacks caught)"],
        ["LANL", "True positive rate at 1 percent false positive rate",
         pct(E1L["tpr_at_1pct_fpr"])],
        ["LANL", "Behavioural only ROC, protocol signal removed", str(E1L["behavioral_only_roc"])],
        ["CIC-IDS2017", "Precision recall area, autoencoder", f"{E1C['autoencoder_prauc']:.3f}"],
        ["CIC-IDS2017", "Precision recall area, Isolation Forest",
         f"{E1C['iforest_prauc']:.3f} (3.1 times random, 4.8 times the rule baseline)"],
        ["CIC-IDS2017", "Precision recall area, random baseline", f"{E1C['random_prauc']:.3f}"],
        ["CIC-IDS2017", "Precision recall area, naive rule baseline",
         f"{E1C['rule_prauc']:.3f} (worse than random)"],
        ["UNSW-NB15", "ROC area under curve", str(E1U["roc_auc"])],
        ["UNSW-NB15", "Precision recall area", f"{E1U['prauc']:.3f}"],
    ], widths=[3.2, 8.4, 5.4], size=9.5)
    B.para("", space_after=8)

    B.rich([("The protocol ablation, and why we ran it. ", True, False),
               ("Every red team login in the dataset used the older NTLM protocol, against about "
                "6 percent of benign logins. That is an extremely powerful signal, and also a "
                "brittle one: it is specific to this dataset and an attacker could simply switch "
                "protocols to evade it. So we deleted the feature entirely and re-ran the "
                f"evaluation. Detection held at {E1L['behavioral_only_roc']}. The result is "
                "driven by generalisable behaviour rather than one fragile artifact. It would "
                "have been easy to keep the higher number quietly.", False, False)])
    B.rich([("The rule baseline is worse than random, and we published it. ", True, False),
               (f"A naive high volume rule scores {E1C['rule_prauc']:.3f} against a random floor of "
                f"{E1C['random_prauc']:.3f}, because stealthy attacks are low volume by definition. "
                "We report it because it is precisely why simple threshold rules fail against "
                "this class of attack.", False, False)])

    B.heading("13.2 Prediction and attribution", 2)
    B.table([
        ["Metric", "Result"],
        ["Markov top 3 accuracy, automatically ordered sequences", pct(E2P["markov_top3"])],
        ["Anti circularity margin over the kill chain baseline", "5.2 times"],
        ["Neural network top 3 accuracy, documented negative result", pct(E2P["lstm_top3"])],
        ["CERT-In verified sequences, top 3 accuracy",
         pct(M["engine2"]["manual_cert_in_top3"])],
        ["Technique embeddings, same tactic cosine similarity",
         f"{E2E['same_tactic_cos']} against {E2E['random_cos']} for random pairs"],
        ["MITRE group profiles used for attribution", "172"],
    ], widths=[10.5, 6.5], size=9.5)

    B.heading("13.3 Operational output on the live campaign", 2)
    B.table([
        ["Output", "Value"],
        ["Events analysed, alerts raised, incidents produced", "2,732, then 1,192, then 1"],
        ["Compromised accounts identified", "104"],
        ["Attack graph size", "479 machines, 502 movements, 4 attacker footholds"],
        ["Critical assets reachable by the attacker", "18"],
        ["Total exposure", "475 machines"],
        ["Effect of isolating the single best choke point", "463 machines severed"],
        ["Events concentrated on the primary foothold", "670 of 702 red team events on C17693"],
    ], widths=[10.5, 6.5], size=9.5)

    B.page_break()

    # ---------------- 14 performance ----------------
    B.heading("14. Performance and scalability", 1)
    B.para("Rather than assert that the system scales, we measured it. The complete pipeline "
              "was timed at nine input sizes on an ordinary laptop processor with no GPU.")
    B.figure(ROOT / "reports" / "scaling_chart_detailed.png",
             "Figure 3. Measured end to end analysis time against input size, with the "
             "measurement table and verification notes.",
             width_cm=17.0)
    B.para("The shipped demonstration campaign completes in 0.140 seconds. The documented "
              "upper limit of 50,000 events in a single analysis completes in 2.493 seconds. "
              "Cost per event stays effectively flat to 20,000 events and rises modestly at the "
              "cap, so we describe the result as fast rather than claiming perfect linearity. "
              "Measurements are the best of three runs after a warm up call, because the first "
              "call loads the model and would otherwise be measuring start up rather than "
              "analysis. Runs above 2,732 events replay the real campaign with offset timestamps, "
              "which is disclosed on the chart itself.")

    # ---------------- 15 honesty ----------------
    B.heading("15. How we kept ourselves honest", 1)
    B.para("Four rules were set at the start of the project and held to, including on the "
              "occasions when they cost us a better looking number.")
    B.bullets([
        [("Never report accuracy. ", True, False),
         ("At this class imbalance, accuracy is a meaningless statistic that flatters a useless "
          "model. We report precision recall area and true positive rate at a fixed false "
          "positive rate, which is what an analyst actually experiences.", False, False)],
        [("Always show baseline lift. ", True, False),
         ("Random and rule baselines were built before the real model, so we could not tune "
          "toward a flattering comparison after the fact. When the rule baseline came out worse "
          "than random, we published that too.", False, False)],
        [("Build the baseline designed to defeat you. ", True, False),
         ("The kill chain order baseline exists purely to test whether our own predictor was "
          "cheating. We report the margin over it, and when the neural model lost to the simple "
          "one, we shipped the simple one.", False, False)],
        [("Nothing fabricated on screen. ", True, False),
         ("Every displayed number traces to the current analysis or to a labelled citation.",
          False, False)],
    ])
    B.para("The fourth rule cost us work, because it meant removing things we had already "
              "built. We list them because a reviewer deserves to know what a team removed, not "
              "only what it added.", space_after=6)
    B.table([
        ["What we removed", "Why it had to go"],
        ["Invented trend lines in the interface",
         "The arrays behind them were decorative fiction. Replaced with real anomaly scores."],
        ["A hard coded detection time claim",
         "It was a fixed string. It is now measured from the timestamps of the log being "
         "analysed, with the industry comparison shown as a citation."],
        ["A fabricated critical asset",
         "The code selected the middle element of a list and presented it as a finding. Replaced "
         "with a stated, defensible heuristic."],
        ["A stale anti circularity figure",
         "The interface claimed a margin our own report contradicted. We found it in a self "
         "audit, corrected it, and then fixed the underlying cause."],
    ], widths=[5.4, 11.6], size=9.5)
    B.para("", space_after=8)
    B.para("That last item produced a structural change rather than a patch. Evaluation "
              "scripts now write a single metrics file, which the interface reads directly. Hand "
              "copied numbers can no longer go stale, because numbers are no longer hand copied.")

    # ---------------- 16 testing ----------------
    B.heading("16. Testing and engineering", 1)
    B.table([
        ["Layer", "Coverage"],
        ["Automated tests", "29 tests covering pipeline correctness, campaign versus per account "
                            "scoping, multi foothold graph reachability, cross screen consistency "
                            "of critical assets, intelligence mapping precision, and resilience "
                            "to a dead feed."],
        ["Browser end to end", "An automated agent drove a real browser through 15 user flows. "
                               "14 passed."],
        ["Deployment", "The container build is verified, and the running container is smoke "
                       "tested including a live external intelligence fetch from inside it."],
    ], widths=[4.2, 12.8], size=9.5)
    B.para("", space_after=8)
    B.para("The browser testing earned its place. It uploaded a file with ISO-8601 timestamps "
              "and found a crash, because every fixture we had written happened to use epoch "
              "integers, and real logs do not. That is a defect a judge would have triggered "
              "during a demonstration. It is fixed with a regression test.")
    B.para("Several tests exist specifically to prevent defects we caused once already: the "
              "vulnerability feed is ordered by vendor rather than by date and must be sorted "
              "explicitly, a ransomware campaign flag must never be mapped as the ransomware "
              "technique itself, and a graph edge must remember every account that used a machine "
              "pair rather than only the first.")

    B.page_break()

    # ---------------- 17 limitations ----------------
    B.heading("17. Limitations we state before you find them", 1)
    B.table([
        ["Limitation", "Our position"],
        ["The demonstration incident carries only three ATT&CK techniques",
         "The dataset is authentication logs only, with no process, file or network telemetry. "
         "Authentication behaviour can honestly evidence pass the hash, brute force and remote "
         "services, and nothing more. We refuse to invent techniques the data cannot support. "
         "Richer telemetry deepens the chain automatically."],
        ["The attribution evaluation scores 100 percent",
         "It is close to trivial by construction, since it retrieves a public profile from part "
         "of itself. We never headline it. The component is transparent retrieval, not a "
         "classifier."],
        ["Prediction on the CERT-In sequences is only 10 percent",
         "That is the honest, non circular figure, and the gap against the automatically ordered "
         "set is itself the finding. Prediction is a supporting feature. The case rests on "
         "detection and correlation."],
        ["Critical assets are a heuristic",
         "The dataset carries no criticality labels. We state the heuristic openly. A real "
         "deployment supplies an asset inventory, which is already an input parameter."],
        ["Containment is simulated",
         "There is no live network to act upon. Every action is labelled simulated and requires "
         "human approval."],
        ["The India scenarios are synthetic",
         "The AIIMS and CBSE scenarios are generated logs styled after real reported incidents, "
         "and are labelled as synthetic in the interface. The LANL campaign is the real data."],
        ["Graph analytics are held in memory",
         "Comfortable to roughly 50,000 events per analysis, as measured. Beyond that the "
         "approach is to shard by tenant or time window, or move to a graph database."],
        ["The application has no user authentication",
         "It is a single analyst demonstration and the login screen is a splash by design. We "
         "make no security claim about the application itself."],
    ], widths=[5.0, 12.0], size=9)

    # ---------------- 18 impact ----------------
    B.heading("18. Impact, deployment path and roadmap", 1)
    B.table([
        ["Dimension", "Situation today", "With Resilience Graph AI"],
        ["Detection", "About 10 days median dwell time",
         "First correlated alert within the log window"],
        ["Analyst load", "1,192 alerts to triage", "1 incident carrying a narrative"],
        ["Containment", "Which of 479 machines do we isolate",
         "Isolate 1 machine, sever 463 machines of exposure"],
        ["Attribution", "Manual reading of threat intelligence",
         "A ranked group with an auditable justification"],
        ["External intelligence", "A separate portal, read separately",
         "Cross referenced against your own techniques"],
    ], widths=[3.4, 6.4, 7.2], size=9.5)
    B.para("", space_after=8)
    B.para("No new sensors are required. The system makes the logs an organisation already "
              "collects tell the story that is already in them. That is what makes it deployable "
              "at the organisations that need it most: hospitals, examination boards, grid "
              "operators, and the large share of government entities running end of life "
              "infrastructure who are least able to fund a 24 hour security operations centre.")
    B.table([
        ["Horizon", "Work"],
        ["30 days", "Connectors for the common log platforms, so logs arrive without manual "
                    "upload."],
        ["90 days", "Operational technology and industrial control coverage, and a graph database "
                    "backend for larger estates."],
        ["6 months", "A sector level view for a national response team, and real containment "
                     "actions placed behind formal change control."],
    ], widths=[3.0, 14.0], size=9.5)

    # ---------------- 19 reproducing ----------------
    B.heading("19. Reproducing this work", 1)
    B.para("The application runs from a fresh clone with no dataset download, because the "
              "trained models, catalogue lookups, embeddings and demonstration scenarios are all "
              "committed to the repository.")
    B.code([
        "# run the application",
        "pip install -r requirements-deploy.txt",
        "python -m uvicorn api.main:app --port 8000",
        "cd frontend && npm install && npm run dev",
        "",
        "# retrain everything from the raw datasets",
        "pip install -r requirements.txt",
        "python -m src.engine1.prep_lanl    && python -m src.engine1.lanl_detect",
        "python -m src.engine1.prep_cicids  && python -m src.engine1.anomaly",
        "python -m src.shared.parse_attack",
        "python -m src.engine2.build_embeddings",
        "python -m src.engine2.build_sequences && python -m src.engine2.build_predictor",
        "python -m scripts.build_cache",
    ])
    B.para("Every evaluation script writes a report to the reports directory. Those reports "
              "are the audit trail behind every number in this document.")

    # ---------------- 20 glossary ----------------
    B.heading("20. Glossary", 1)
    B.table([
        ["Term", "Meaning"],
        ["ATT&CK", "MITRE's public catalogue of attacker techniques, each with an identifier "
                   "such as T1110."],
        ["Blast radius", "Everything an attacker can reach from where they currently are."],
        ["CERT-In", "India's national Computer Emergency Response Team."],
        ["Choke point", "A machine that many attack paths pass through, and therefore the best "
                        "one to isolate."],
        ["Critical asset, or crown jewel", "A machine worth protecting most, such as a patient "
                                           "database or a domain controller."],
        ["Dwell time", "How long an attacker remains undetected inside a network."],
        ["False positive rate", "The share of normal events wrongly flagged. This is what causes "
                                "alert fatigue."],
        ["Isolation Forest", "An algorithm that finds outliers by measuring how easily a point "
                             "can be separated from the rest."],
        ["Lateral movement", "Moving from machine to machine inside a network after gaining "
                             "entry."],
        ["NTLM", "An older Windows authentication protocol, vulnerable to pass the hash."],
        ["Pass the hash", "Logging in using a stolen credential fingerprint instead of the "
                          "password itself."],
        ["PR-AUC", "Precision recall area under curve. The correct substitute for accuracy when "
                   "the event of interest is rare."],
        ["Red team", "Authorised attackers who simulate a real intrusion. Their recorded actions "
                     "are the ground truth labels."],
        ["ROC-AUC", "The probability that the model ranks a random attack above a random benign "
                    "event. 1.0 is perfect, 0.5 is a coin toss."],
        ["SIEM", "The log collection platform a security team runs on."],
        ["SOAR", "Security orchestration, automation and response. The layer that acts after "
                 "detection."],
        ["SOC", "Security Operations Centre. The team watching for attacks."],
        ["TPR", "True positive rate, also called recall. The share of real attacks caught."],
        ["Unsupervised learning", "Training without labels. The model learns what normal looks "
                                  "like and flags deviation from it."],
    ], widths=[4.4, 12.6], size=9)

    B.para("", space_after=12)
    B.para("Resilience Graph AI, Team rishiikumarsingh2201, ET AI Hackathon 2026, Problem "
              "Statement 7.", size=9, italic=True, align="center")



# ==========================================================================
# self check: the brief bans em dashes, en dashes and emoji outright
# ==========================================================================
FORBIDDEN = {
    "—": "em dash",
    "–": "en dash",
    "‒": "figure dash",
    "―": "horizontal bar",
}
EMOJI = re.compile(
    "[\U0001F000-\U0001FAFF"      # pictographs, emoticons, symbols
    "☀-➿"               # misc symbols and dingbats
    "️"                      # variation selector
    "←-⇿"               # arrows
    "⬀-⯿]"              # misc symbols and arrows
)


def self_check(texts) -> None:
    """Fail loudly if a banned character reached the document."""
    problems = []
    for text in texts:
        t = str(text)
        for ch, name in FORBIDDEN.items():
            if ch in t:
                problems.append("%s in: %s" % (name, t[:70]))
        m = EMOJI.search(t)
        if m:
            problems.append("emoji or symbol %r in: %s" % (m.group(), t[:70]))
    if problems:
        raise SystemExit("Banned characters found:\n  " + "\n  ".join(problems))
    print("self-check passed: %d text blocks, no em dash, en dash or emoji" % len(texts))


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    d = DocxBackend()
    build(d)
    self_check(d.texts)
    d.save(OUT_DOCX)
    print("wrote %s" % OUT_DOCX.relative_to(ROOT))

    h = HtmlBackend()
    build(h)
    h.save(OUT_HTML)

    subprocess.run(
        [str(CHROME), "--headless=new", "--disable-gpu", "--no-pdf-header-footer",
         "--print-to-pdf=%s" % OUT_PDF, OUT_HTML.resolve().as_uri()],
        check=True, capture_output=True, timeout=300)
    print("wrote %s" % OUT_PDF.relative_to(ROOT))


if __name__ == "__main__":
    main()
